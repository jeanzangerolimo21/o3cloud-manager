import csv
import io
import uuid
from pathlib import Path
from datetime import date, datetime
from decimal import Decimal

from docx import Document
from docx.shared import Inches
from openpyxl import Workbook
from openpyxl.drawing.image import Image as XlsxImage
from openpyxl.styles import Font, PatternFill
from reportlab.lib import colors
from reportlab.lib.pagesizes import A4, landscape
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.platypus import Image, Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle

from app.core.email import EmailService
from app.core.filters import cnpj_br, date_br, datetime_br, moeda
from app.core.storage import StorageService
from app.relatorios.catalogo import TIPO_DATA, TIPO_DATETIME, TIPO_MOEDA, TIPOS_NUMERICOS, buscar_fonte, listar_fontes
from app.relatorios.repository import RelatorioRepository


LOGO_PATH = "/opt/o3cloud-manager/app/static/img/logo.png"
PREVIEW_LIMIT = 500
EXPORT_SYNC_LIMIT = 5000
BACKGROUND_LIMIT = 100000
FONTES_PERIODO_OBRIGATORIO = {"contratos", "inadimplencias", "faturamentos", "leads", "oportunidades", "implantacoes", "demandas"}
AGREGACOES = {"COUNT": "CONTAGEM", "SUM": "SOMA", "AVG": "MEDIA", "MIN": "MINIMO", "MAX": "MAXIMO", "COUNT_DISTINCT": "CONTAGEM DISTINTA"}


class RelatorioService:
    repository = RelatorioRepository

    @classmethod
    def contexto(cls, usuario_id=None, perfil_codigo=None):
        return {
            "fontes": listar_fontes(),
            "modelos": cls.repository.listar_modelos(usuario_id, perfil_codigo),
            "jobs": cls.repository.listar_jobs_usuario(usuario_id),
            "agregacoes": AGREGACOES,
        }

    @classmethod
    def buscar_modelo(cls, modelo_id):
        return cls.repository.buscar_modelo(modelo_id)

    @classmethod
    def executar(cls, dados, usuario=None, modelo_id=None, limite=500, formato="HTML"):
        config = cls._normalizar_config(dados)
        fonte = buscar_fonte(config["fonte"])
        if not fonte:
            raise ValueError("Fonte de relatório inválida.")
        campos = cls._campos_selecionados(fonte, config["campos"])
        if not campos and not config["agregacoes"]:
            raise ValueError("Selecione ao menos um campo ou cálculo.")

        select_parts = []
        colunas = []
        group_by = []
        agrupamentos_config = set(config["agrupamentos"] or [])
        for campo in campos:
            select_parts.append(f"{campo.expressao} AS `{campo.codigo}`")
            colunas.append(campo.as_dict())
            if campo.agrupavel and (campo.codigo in agrupamentos_config or (config["agregacoes"] and not agrupamentos_config)):
                group_by.append(campo.expressao)
        for codigo in agrupamentos_config:
            campo = fonte.campo(codigo)
            if campo and campo.agrupavel and campo.expressao not in group_by:
                group_by.append(campo.expressao)

        for item in config["agregacoes"]:
            campo = fonte.campo(item.get("campo"))
            funcao = item.get("funcao")
            if not campo or not campo.agregavel or funcao not in AGREGACOES:
                continue
            alias = f"{funcao.lower()}_{campo.codigo}"
            expr = "COUNT(*)" if funcao == "COUNT" else f"{funcao.replace('_DISTINCT', '')}({('DISTINCT ' if funcao == 'COUNT_DISTINCT' else '')}{campo.expressao})"
            select_parts.append(f"{expr} AS `{alias}`")
            colunas.append({"codigo": alias, "nome": f"{AGREGACOES[funcao]} - {campo.nome}", "tipo": campo.tipo if funcao != "COUNT" else "INTEIRO", "formato": campo.formato})

        where, params = cls._where(fonte, config["filtros"])
        sql = "SELECT " + ", ".join(select_parts) + "\n" + fonte.from_sql
        if where:
            sql += "\nWHERE " + " AND ".join(where)
        if group_by:
            sql += "\nGROUP BY " + ", ".join(group_by)

        ordem = cls._order_by(fonte, config["ordenacao"])
        if ordem:
            sql += "\nORDER BY " + ordem
        limite_normalizado = cls._limite(limite)
        sql += "\nLIMIT %s"
        params.append(limite_normalizado + 1)

        linhas = cls.repository.executar_sql(sql, params)
        truncado = len(linhas) > limite_normalizado
        if truncado:
            linhas = linhas[:limite_normalizado]
        cls.repository.auditar_execucao({
            "modelo_id": modelo_id,
            "fonte": fonte.codigo,
            "formato": formato,
            "total_linhas": len(linhas),
            "usuario_id": (usuario or {}).get("id"),
            "usuario_email": (usuario or {}).get("email"),
            "filtros": config["filtros"],
        })
        linhas_formatadas = [
            {col["codigo"]: cls._formatar(linha.get(col["codigo"]), col) for col in colunas}
            for linha in linhas
        ]
        return {"fonte": fonte.as_dict(), "config": config, "colunas": colunas, "linhas": linhas, "linhas_formatadas": linhas_formatadas, "total": len(linhas), "limite": limite_normalizado, "truncado": truncado}


    @classmethod
    def validar_execucao_sincrona(cls, dados, limite=PREVIEW_LIMIT, exigir_periodo=False):
        config = cls._normalizar_config(dados)
        fonte = buscar_fonte(config["fonte"])
        if not fonte:
            raise ValueError("Fonte de relatório inválida.")
        if cls._limite(limite) > PREVIEW_LIMIT and fonte.codigo in FONTES_PERIODO_OBRIGATORIO and not cls._tem_filtro_periodo(fonte, config["filtros"]):
            raise ValueError("Para exportações diretas desta fonte, informe um filtro de período. Para relatório completo, use execução em segundo plano.")
        if exigir_periodo and fonte.codigo in FONTES_PERIODO_OBRIGATORIO and not cls._tem_filtro_periodo(fonte, config["filtros"]):
            raise ValueError("Informe um filtro de período ou solicite a execução em segundo plano.")
        return config

    @classmethod
    def solicitar_job(cls, dados, usuario, modelo_id=None, formato="XLSX"):
        if not (usuario or {}).get("email"):
            raise ValueError("Usuário logado sem e-mail para envio do relatório.")
        config = cls._normalizar_config(dados)
        fonte = buscar_fonte(config["fonte"])
        if not fonte:
            raise ValueError("Fonte de relatório inválida.")
        return cls.repository.inserir_job({
            "modelo_id": modelo_id,
            "fonte": fonte.codigo,
            "formato": (formato or "XLSX").upper(),
            "configuracao": config,
            "usuario_id": usuario.get("id"),
            "usuario_email": usuario.get("email"),
        })

    @classmethod
    def processar_jobs(cls, limite=1):
        processados = []
        for _ in range(max(1, int(limite or 1))):
            job = cls.repository.proximo_job_pendente()
            if not job:
                break
            if not cls.repository.marcar_job_processando(job["id"]):
                continue
            try:
                processados.append(cls._processar_job(job))
            except Exception as erro:
                cls.repository.falhar_job(job["id"], erro)
                processados.append({"job_id": job["id"], "status": "ERRO", "erro": str(erro)})
        return processados

    @classmethod
    def _processar_job(cls, job):
        config = cls.repository._json(job.get("configuracao_json"), {})
        class Dados(dict):
            def getlist(self, key):
                valor = self.get(key, [])
                return valor if isinstance(valor, list) else [valor]
        dados = Dados(config)
        formato = (job.get("formato") or "XLSX").lower()
        resultado = cls.executar(dados, {"id": job.get("solicitado_por_id"), "email": job.get("solicitado_por_email")}, job.get("modelo_id"), BACKGROUND_LIMIT, formato.upper())
        conteudo = cls._exportar_por_formato(resultado, formato, job.get("solicitado_por_email"))
        nome = f"relatorio-{job['id']}-{uuid.uuid4().hex}.{formato}"
        pasta = StorageService.BASE_STORAGE / "relatorios"
        pasta.mkdir(parents=True, exist_ok=True)
        caminho = pasta / nome
        caminho.write_bytes(conteudo)
        url = StorageService.url("relatorios", nome)
        email_resultado = cls._enviar_email_job(job, url, resultado)
        cls.repository.concluir_job(job["id"], {
            "total_linhas": resultado["total"],
            "arquivo_nome": nome,
            "arquivo_url": url,
            "email_enviado": email_resultado.get("enviado"),
            "email_erro": None if email_resultado.get("enviado") else email_resultado.get("motivo"),
        })
        return {"job_id": job["id"], "status": "CONCLUIDO", "arquivo_url": url, "email": email_resultado}

    @classmethod
    def _exportar_por_formato(cls, resultado, formato, usuario_email=""):
        if formato == "csv":
            return cls.exportar_csv(resultado)
        if formato == "docx":
            return cls.exportar_docx(resultado, usuario_email)
        if formato == "pdf":
            return cls.exportar_pdf(resultado, usuario_email)
        return cls.exportar_xlsx(resultado, usuario_email)

    @classmethod
    def _enviar_email_job(cls, job, url, resultado):
        assunto = "Relatório O3Cloud disponível para download"
        corpo = (
            "Seu relatório foi concluído.\n\n"
            f"Fonte: {resultado['fonte']['nome']}\n"
            f"Total de linhas: {resultado['total']}\n"
            f"Link para download: {url}\n\n"
            "O link fica disponível conforme a política de armazenamento do O3Cloud Manager."
        )
        try:
            return EmailService.enviar(assunto, corpo, [job.get("solicitado_por_email")])
        except Exception as erro:
            return {"enviado": False, "motivo": str(erro)}

    @classmethod
    def _tem_filtro_periodo(cls, fonte, filtros):
        for filtro in filtros:
            campo = fonte.campo(filtro.get("campo"))
            if campo and campo.tipo in (TIPO_DATA, TIPO_DATETIME) and filtro.get("valor"):
                return True
        return False

    @staticmethod
    def _limite(limite):
        try:
            return min(max(int(limite or PREVIEW_LIMIT), 1), BACKGROUND_LIMIT)
        except (TypeError, ValueError):
            return PREVIEW_LIMIT

    @classmethod
    def salvar_modelo(cls, dados, usuario, modelo_id=None):
        config = cls._normalizar_config(dados)
        nome = cls._texto(dados.get("nome"))
        if not nome:
            raise ValueError("Nome do relatório é obrigatório.")
        visibilidade = cls._texto(dados.get("visibilidade")) or "PRIVADO"
        if visibilidade not in ("PRIVADO", "PERFIL", "GLOBAL"):
            visibilidade = "PRIVADO"
        payload = {
            "nome": nome,
            "descricao": cls._texto(dados.get("descricao")),
            "fonte": config["fonte"],
            "configuracao": config,
            "visibilidade": visibilidade,
            "perfis": [p.strip().upper() for p in (dados.get("perfis") or "").split(",") if p.strip()],
            "usuario_id": usuario.get("id"),
            "usuario_email": usuario.get("email"),
        }
        if modelo_id:
            cls.repository.atualizar_modelo(modelo_id, payload)
            return modelo_id
        return cls.repository.inserir_modelo(payload)

    @classmethod
    def exportar_csv(cls, resultado):
        buffer = io.StringIO()
        writer = csv.writer(buffer, delimiter=";")
        writer.writerow(["O3Cloud Manager", resultado["fonte"]["nome"], datetime.now().strftime("%d/%m/%Y %H:%M")])
        writer.writerow([col["nome"] for col in resultado["colunas"]])
        for linha in resultado["linhas"]:
            writer.writerow([cls._formatar(linha.get(col["codigo"]), col) for col in resultado["colunas"]])
        return buffer.getvalue().encode("utf-8-sig")

    @classmethod
    def exportar_xlsx(cls, resultado, usuario_email=""):
        wb = Workbook()
        ws = wb.active
        ws.title = "Relatorio"
        try:
            img = XlsxImage(LOGO_PATH)
            img.width = 140
            img.height = 50
            ws.add_image(img, "A1")
        except Exception:
            ws["A1"] = "O3Cloud Manager"
        ws["A4"] = resultado["fonte"]["nome"]
        ws["A4"].font = Font(bold=True, size=14)
        ws["A5"] = f"Gerado em {datetime.now().strftime('%d/%m/%Y %H:%M')} por {usuario_email or 'sistema'}"
        header_row = 7
        for idx, col in enumerate(resultado["colunas"], 1):
            cell = ws.cell(header_row, idx, col["nome"])
            cell.font = Font(bold=True, color="FFFFFF")
            cell.fill = PatternFill("solid", fgColor="1F4E78")
        for row_idx, linha in enumerate(resultado["linhas"], header_row + 1):
            for col_idx, col in enumerate(resultado["colunas"], 1):
                ws.cell(row_idx, col_idx, cls._valor_planilha(linha.get(col["codigo"])))
        for column_cells in ws.columns:
            ws.column_dimensions[column_cells[0].column_letter].width = min(max(len(str(column_cells[0].value or "")) + 4, 14), 40)
        buffer = io.BytesIO()
        wb.save(buffer)
        return buffer.getvalue()

    @classmethod
    def exportar_docx(cls, resultado, usuario_email=""):
        doc = Document()
        try:
            doc.add_picture(LOGO_PATH, width=Inches(1.7))
        except Exception:
            doc.add_paragraph("O3Cloud Manager")
        doc.add_heading(resultado["fonte"]["nome"], level=1)
        doc.add_paragraph(f"Gerado em {datetime.now().strftime('%d/%m/%Y %H:%M')} por {usuario_email or 'sistema'}")
        tabela = doc.add_table(rows=1, cols=len(resultado["colunas"]))
        tabela.style = "Table Grid"
        for idx, col in enumerate(resultado["colunas"]):
            tabela.rows[0].cells[idx].text = col["nome"]
        for linha in resultado["linhas"]:
            row = tabela.add_row().cells
            for idx, col in enumerate(resultado["colunas"]):
                row[idx].text = cls._formatar(linha.get(col["codigo"]), col)
        buffer = io.BytesIO()
        doc.save(buffer)
        return buffer.getvalue()

    @classmethod
    def exportar_pdf(cls, resultado, usuario_email=""):
        buffer = io.BytesIO()
        doc = SimpleDocTemplate(buffer, pagesize=landscape(A4), rightMargin=24, leftMargin=24, topMargin=24, bottomMargin=24)
        styles = getSampleStyleSheet()
        elements = []
        try:
            elements.append(Image(LOGO_PATH, width=100, height=42))
        except Exception:
            elements.append(Paragraph("O3Cloud Manager", styles["Title"]))
        elements.append(Paragraph(resultado["fonte"]["nome"], styles["Heading1"]))
        elements.append(Paragraph(f"Gerado em {datetime.now().strftime('%d/%m/%Y %H:%M')} por {usuario_email or 'sistema'}", styles["Normal"]))
        elements.append(Spacer(1, 12))
        data = [[col["nome"] for col in resultado["colunas"]]]
        for linha in resultado["linhas"][:300]:
            data.append([cls._formatar(linha.get(col["codigo"]), col) for col in resultado["colunas"]])
        table = Table(data, repeatRows=1)
        table.setStyle(TableStyle([
            ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#1F4E78")),
            ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
            ("GRID", (0, 0), (-1, -1), 0.25, colors.grey),
            ("FONTSIZE", (0, 0), (-1, -1), 8),
        ]))
        elements.append(table)
        doc.build(elements)
        return buffer.getvalue()

    @classmethod
    def _normalizar_config(cls, dados):
        return {
            "fonte": cls._texto(dados.get("fonte")),
            "campos": cls._campos_form(dados),
            "filtros": cls._filtros_form(dados),
            "ordenacao": {"campo": cls._texto(dados.get("ordem_campo")), "direcao": cls._texto(dados.get("ordem_direcao")) or "ASC"},
            "agrupamentos": cls._lista(dados.getlist("agrupamentos") if hasattr(dados, "getlist") else dados.get("agrupamentos")),
            "agregacoes": cls._agregacoes_form(dados),
        }

    @classmethod
    def _campos_selecionados(cls, fonte, codigos):
        campos = []
        for codigo in codigos:
            campo = fonte.campo(codigo)
            if campo and campo.codigo not in [c.codigo for c in campos]:
                campos.append(campo)
        return campos

    @classmethod
    def _campos_form(cls, dados):
        getlist = dados.getlist if hasattr(dados, "getlist") else lambda key: dados.get(key, [])
        campos = cls._lista(getlist("campos"))

        def ordem(codigo):
            try:
                return int(dados.get(f"campo_ordem_{codigo}") or 999)
            except (TypeError, ValueError):
                return 999

        return sorted(campos, key=lambda codigo: (ordem(codigo), campos.index(codigo)))

    @classmethod
    def _where(cls, fonte, filtros):
        where = [fonte.where_base] if fonte.where_base else []
        params = []
        for filtro in filtros:
            campo = fonte.campo(filtro.get("campo"))
            operador = filtro.get("operador")
            valor = filtro.get("valor")
            valor_final = filtro.get("valor_final")
            if not campo or not campo.filtravel or valor in (None, ""):
                continue
            expr = campo.expressao
            if operador == "igual":
                where.append(f"{expr} = %s")
                params.append(valor)
            elif operador == "diferente":
                where.append(f"{expr} <> %s")
                params.append(valor)
            elif operador == "contem":
                where.append(f"{expr} LIKE %s")
                params.append(f"%{valor}%")
            elif operador == "comeca":
                where.append(f"{expr} LIKE %s")
                params.append(f"{valor}%")
            elif operador in ("maior", "maior_igual", "menor", "menor_igual") and campo.tipo in TIPOS_NUMERICOS | {TIPO_DATA, TIPO_DATETIME}:
                sinal = {"maior": ">", "maior_igual": ">=", "menor": "<", "menor_igual": "<="}[operador]
                where.append(f"{expr} {sinal} %s")
                params.append(valor)
            elif operador == "entre" and valor_final:
                where.append(f"{expr} BETWEEN %s AND %s")
                params.extend([valor, valor_final])
        return where, params

    @classmethod
    def _order_by(cls, fonte, ordenacao):
        campo = fonte.campo((ordenacao or {}).get("campo"))
        if not campo or not campo.ordenavel:
            return ""
        direcao = "DESC" if (ordenacao or {}).get("direcao") == "DESC" else "ASC"
        return f"{campo.expressao} {direcao}"

    @classmethod
    def _filtros_form(cls, dados):
        filtros = []
        getlist = dados.getlist if hasattr(dados, "getlist") else lambda key: dados.get(key, [])
        campos = getlist("filtro_campo")
        operadores = getlist("filtro_operador")
        valores = getlist("filtro_valor")
        finais = getlist("filtro_valor_final")
        for idx, campo in enumerate(campos):
            filtros.append({
                "campo": cls._texto(campo),
                "operador": cls._texto(operadores[idx] if idx < len(operadores) else ""),
                "valor": cls._texto(valores[idx] if idx < len(valores) else ""),
                "valor_final": cls._texto(finais[idx] if idx < len(finais) else ""),
            })
        return filtros

    @classmethod
    def _agregacoes_form(cls, dados):
        agregacoes = []
        getlist = dados.getlist if hasattr(dados, "getlist") else lambda key: dados.get(key, [])
        campos = getlist("agregacao_campo")
        funcoes = getlist("agregacao_funcao")
        for idx, campo in enumerate(campos):
            if campo:
                agregacoes.append({"campo": cls._texto(campo), "funcao": cls._texto(funcoes[idx] if idx < len(funcoes) else "")})
        return agregacoes

    @staticmethod
    def _lista(valor):
        if not valor:
            return []
        if isinstance(valor, str):
            return [valor]
        return [str(item) for item in valor if item]

    @staticmethod
    def _texto(valor):
        return str(valor).strip() if valor is not None else ""

    @classmethod
    def _formatar(cls, valor, col):
        if valor is None:
            return ""
        if col.get("formato") == "moeda" or col.get("tipo") == TIPO_MOEDA:
            return moeda(valor)
        if col.get("formato") == "cnpj":
            return cnpj_br(valor)
        if isinstance(valor, datetime):
            return datetime_br(valor)
        if isinstance(valor, date):
            return date_br(valor)
        return str(valor)

    @staticmethod
    def _valor_planilha(valor):
        if isinstance(valor, Decimal):
            return float(valor)
        return valor
